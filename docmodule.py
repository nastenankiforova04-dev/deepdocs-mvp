import io
from pprint import pprint
from typing import Any

import pandas as pd
from docx import Document
import re
import pdfplumber
import gspread  # Добавлен gspread


def extract_data(file):
    file_name_str = getattr(file, "name", "").lower()

    party_regex = r"((?:ООО|ИП|АО|ПАО|ЗАО|ГБУ|Общество с ограниченной ответственностью|Акционерное общество)\s+[«\"'][^»\"']+[»\"']|[A-Za-z\s&]+(?:A/S|Ltd\.?|LLC|Inc\.?|GmbH))"
    number_regex = r"(?:Договор[а-я]{0,3}|№)\s*[№N ]?\s*([а-яА-ЯёЁ]*\d[\w\-/№]*)"
    date_regex = r"(\d{1,2}[.\s](?:\d{2}|[а-яА-ЯёЁ]{3,10})[.\s]\d{2,4})"
    summ_regex = r"(?:стоимость|сумма|сумму|цена|всего)\D{0,100}?(\d[\d\s,.]*?\d)\s*(?:\([^)]+\)\s*)?(?:руб|₽|usd|евро)"

    results = {
        "number": "Не найден",
        "dates": [],
        "all_parties": [],
        "conditions": [],
        "sums": []
    }

    def extract_data_from_text(text: str):
        text = text.strip()
        if not text:
            return

        match_sums = re.findall(summ_regex, text, re.IGNORECASE | re.DOTALL)
        for summ in match_sums:
            results["sums"].append(summ.strip())

        found_dates = re.findall(date_regex, text)
        results["dates"].extend([d.strip() for d in found_dates[::-1]])

        parties = re.findall(party_regex, text)
        for p_party in parties:
            if p_party.strip() not in results["all_parties"]:
                results["all_parties"].append(p_party.strip())

        if results["number"] == "Не найден":
            match_no = re.search(number_regex, text, re.I)
            if match_no:
                results["number"] = match_no.group(1)

        if re.search(r'[а-яА-ЯёЁ]', text):
            results["conditions"].append(text)

    # === ЛОГИКА ДЛЯ DOCX ===
    if file_name_str.endswith('.docx'):
        doc = Document(file)
        paragraphs = doc.paragraphs

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)

        for paragraph in paragraphs:
            extract_data_from_text(paragraph.text)

    # === ЛОГИКА ДЛЯ PDF ===
    elif file_name_str.endswith('.pdf'):
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        extract_data_from_text(line)

    results["date"] = results["dates"][0] if results["dates"] else ""
    results["agent"] = results["all_parties"][:1]
    results["counterparty"] = results["all_parties"][1:2]

    pattern_start = r"^(?:[а-яА-ЯёЁ\w]+\s+)?1[\.\s](?!\d)"
    pattern_end_3 = r"^(?:[а-яА-ЯёЁ\w]+\s+)?3[\.\s](?!\d)"
    pattern_end_2 = r"^(?:[а-яА-ЯёЁ\w]+\s+)?2[\.\s](?!\d)"

    def count_subpoints(text):
        match = re.match(r"^(?:[а-яА-ЯёЁ\w]+\s+)?(\d+(?:\.\d+)*)", text.strip())
        if match:
            num_part = match.group(1)
            digits = re.findall(r"\d+", num_part)
            return len(digits)
        return 0

    try:
        non_repetitive_conditions = []
        for i in range(len(results["conditions"]) - 1):
            if count_subpoints(results["conditions"][i]) == 1 and count_subpoints(results["conditions"][i + 1]) == 1:
                continue
            non_repetitive_conditions.append(results["conditions"][i])

        index_start = next(idx for idx, text in enumerate(non_repetitive_conditions) if re.search(pattern_start, text))

        try:
            index_end = next(idx for idx, text in enumerate(non_repetitive_conditions[index_start + 1:]) if
                             re.search(pattern_end_3, text)) + index_start + 1
        except StopIteration:
            try:
                index_end = next(idx for idx, text in enumerate(non_repetitive_conditions[index_start + 1:]) if
                                 re.search(pattern_end_2, text)) + index_start + 1
            except StopIteration:
                index_end = len(non_repetitive_conditions)

        results["Предмет_договора"] = non_repetitive_conditions[index_start + 1:index_end]

    except StopIteration:
        results["Предмет_договора"] = ["Не найдены"]

    def clear_summ(summ: str):
        clean_amount = re.sub(r'[^\d,.]', '', summ)
        if ',' in clean_amount:
            clean_amount = clean_amount.replace(',', '.')
        try:
            return float(clean_amount)
        except ValueError:
            return 0

    results["summ"] = max(map(clear_summ, results["sums"])) if results["sums"] else 0

    return results


def json_to_excel_bytes(json_data: list[dict[str, Any]]):
    df = pd.DataFrame(json_data)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Договоры')
    return output.getvalue()


def upload_to_google_sheets(json_data: list[dict[str, Any]], sheet_id: str, credits_path: str = "credentials.json"):
    """
    Открывает Google Таблицу по ID и дописывает новые договоры вниз.
    """
    import gspread

    gc = gspread.service_account(filename=credits_path)

    sh = gc.open_by_key(sheet_id)
    worksheet = sh.get_worksheet(0)

    df = pd.DataFrame(json_data)
    rows = df.values.tolist()

    cleaned_rows = [[str(cell) if pd.notna(cell) else "" for cell in row] for row in rows]

    existing_data = worksheet.get_all_values()
    if not existing_data:
        # Если таблица абсолютно пустая - сначала вставляем названия колонок
        header = df.columns.tolist()
        worksheet.append_row(header)

    worksheet.append_rows(cleaned_rows)

    return True